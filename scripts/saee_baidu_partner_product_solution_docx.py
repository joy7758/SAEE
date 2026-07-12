#!/usr/bin/env python3
"""Build the human-fillable Baidu partner product solution DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "cloud-entry-package/materials/SAEE_BAIDU_PARTNER_PRODUCT_SOLUTION_V1.docx"
ARCHITECTURE = ROOT / "cloud-entry-package/architecture.png"
NAVY = RGBColor(23, 32, 51)
BLUE = RGBColor(47, 111, 235)
MUTED = RGBColor(83, 96, 120)
RED = RGBColor(155, 28, 28)


def set_font(run, *, size: float, bold: bool = False, color: RGBColor = NAVY) -> None:
    # Named Unicode render override: LibreOffice must resolve the same family
    # for Latin and CJK runs during the mandatory visual QA conversion.
    run.font.name = "Arial Unicode MS"
    rfonts = run._element.get_or_add_rPr().rFonts
    rfonts.set(qn("w:ascii"), "Arial Unicode MS")
    rfonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    rfonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def set_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("SAEE × 百度智能云 | 伙伴产品方案包"), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("partner consultation submitted · marketplace submission=false"), size=8, color=MUTED)


def add_para(doc: Document, text: str, *, bold_prefix: str | None = None, color: RGBColor = NAVY) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), size=11, bold=True, color=color)
        set_font(p.add_run(text[len(bold_prefix):]), size=11, color=color)
    else:
        set_font(p.add_run(text), size=11, color=color)


def add_field(doc: Document, label: str, value: str = "TBD_OWNER_INPUT") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(f"{label}："), size=10.5, bold=True)
    set_font(p.add_run(value), size=10.5, color=RED if value.startswith("TBD") else MUTED)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def build() -> Path:
    doc = Document()
    style_document(doc)
    set_header_footer(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run("百度智能云伙伴产品方案"), size=30, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    set_font(p.add_run("SAEE Agent Readiness Platform / SAEE 智能体上线准备平台"), size=16, color=BLUE)
    add_para(doc, "合作方向：千帆 Agent 上线前执行证据评估与应用场景共建")
    add_para(doc, "文档状态：千帆伙伴咨询已提交；更高阶合作资质字段仍待负责人补全")
    add_para(doc, "版本：v1.0 · 2026-07-13", color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    set_font(p.add_run("重要边界"), size=12, bold=True, color=RED)
    add_para(doc, "本方案不声明百度官方集成、产品认证、云市场入驻、客户验证或生产就绪。", color=RED)

    page_break(doc)
    doc.add_heading("1. 后续合作主体信息（负责人必填）", level=1)
    add_field(doc, "法定主体中文名")
    add_field(doc, "统一社会信用代码")
    add_field(doc, "单位性质")
    add_field(doc, "单位官网")
    add_field(doc, "单位地址")
    add_field(doc, "联系人姓名")
    add_field(doc, "联系人职位")
    add_field(doc, "联系人手机号")
    add_field(doc, "联系人邮箱")
    add_field(doc, "支持热线及服务时间")
    doc.add_heading("建议表单选择", level=2)
    add_para(doc, "产品形态：服务；是否同时选择 SaaS 由负责人根据真实交付形态决定。")
    add_para(doc, "服务能力：产品集成。合作权益：应用场景共建、技术赋能提升。")
    add_para(doc, "不得将未核验字段自动补全或从个人资料中推断。", color=RED)

    page_break(doc)
    doc.add_heading("2. 产品介绍", level=1)
    add_para(doc, "SAEE 是用于评估 AI Agent 在真实部署前是否具备充分执行证据的智能体就绪基础设施。")
    add_para(doc, "首个产品是 SAEE Agent Readiness Assessment：客户提供 Agent 配置、Tool 调用记录、Execution Trace 与 Evidence Bundle，系统返回证据覆盖、缺失证据、风险信号和有边界的上线准备建议。")
    doc.add_heading("两个公共操作", level=2)
    add_para(doc, "saee.evaluate_agent_run：评估一次 Agent 运行的执行证据覆盖与缺口。")
    add_para(doc, "saee.evaluate_evidence：评估 Evidence Bundle 的质量、覆盖和限制。")
    doc.add_heading("客户价值", level=2)
    add_para(doc, "把“这次跑通”与“具备上线所需证据”分开，帮助企业在退款、支付、代码发布等高影响场景中识别缺少的回滚、权限和人工审批证据。")
    add_para(doc, "评估结果不授权部署、付款、权限扩大或任何外部动作。")

    page_break(doc)
    doc.add_heading("3. 百度智能云目标组合架构", level=1)
    doc.add_picture(str(ARCHITECTURE), width=Inches(6.45))
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "目标路径：BOS 脱敏对象引用 → 百度千帆 → Agent 应用 → SAEE Connector → Evaluation Engine → Evidence Analysis → Readiness Report。")
    add_para(doc, "当前 Alpha 已完成两个只读 MCP 工具、Qianfan-style 离线 host simulation，以及两个合成业务场景的真实 Qianfan 产品 roundtrip；尚未访问 BOS，也不构成官方千帆集成。")

    page_break(doc)
    doc.add_heading("4. Demo、交付与合作建议", level=1)
    doc.add_heading("本地可核验 Demo", level=2)
    add_para(doc, "智能客服退款 Agent：证据覆盖 75%，缺失 HUMAN_APPROVAL，建议 HUMAN_REVIEW_REQUIRED；不执行退款。")
    add_para(doc, "代码发布 Agent：证据覆盖 50%，缺失 ROLLBACK_PLAN 与 HUMAN_APPROVAL，建议 REPLAN；不执行部署。")
    add_para(doc, "证据包评估：检查 evidence quality；不认证 trace 真实性。")
    doc.add_heading("本地材料", level=2)
    add_para(doc, "30 分钟 Cloud Entry Package、10 页技术白皮书、3 分钟 Demo 视频、OpenAPI、MCP、Capability Card、架构图、截图与离线 validator 已准备。")
    doc.add_heading("建议合作路径", level=2)
    add_para(doc, "第一步：千帆伙伴咨询已按授权提交，方向为产品集成与应用场景共建；等待并记录百度反馈。")
    add_para(doc, "第二步：两个真实 Qianfan 合成场景 roundtrip 已完成；下一项是百度工程审阅，不把 provider 调用升级为官方集成。")
    add_para(doc, "第三步：在资质、软著、支持、企业账号和协议条件满足后，再评估产品认证或云市场入驻。")
    doc.add_heading("当前真值", level=2)
    add_para(doc, "official_qianfan_integration=false · customer_validated=false · marketplace_submission=false · production_ready=false", color=RED)

    page_break(doc)
    doc.add_heading("5. 官方入口与附件清单", level=1)
    add_para(doc, "千帆伙伴咨询：https://cloud.baidu.com/survey/qianfanpartnerconsultation.html?track=C841333")
    add_para(doc, "合作伙伴申请：https://cloud.baidu.com/partner/apply.html")
    add_para(doc, "产品认证：https://cloud.baidu.com/partner/product-certification.html")
    add_para(doc, "云市场条件：https://cloud.baidu.com/doc/Market/s/ojy6wl8sd")
    add_para(doc, "云市场流程：https://cloud.baidu.com/doc/Market/s/9jy6y1c8f")
    doc.add_heading("建议随附文件（人工确认后）", level=2)
    add_para(doc, "本 Word 产品方案、技术白皮书 PDF、Demo 视频或受控地址、GitHub Release 地址、营业执照/资质材料。")
    add_para(doc, "千帆伙伴咨询已按授权提交；后续上传附件、接受协议或向其他入口复用联系人数据仍需单独授权。", color=RED)

    props = doc.core_properties
    props.title = "SAEE 百度智能云伙伴产品方案 v1.0"
    props.subject = "Qianfan partner consultation submitted; broader partner qualification handoff"
    props.author = "SAEE"
    props.keywords = "SAEE, Baidu Qianfan, Agent Readiness, local draft"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
