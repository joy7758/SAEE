#!/usr/bin/env python3
"""Fail-closed checks for the local SAEE software-copyright application pack."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
PACK = ROOT / "docs/ip/software-copyright"
MANIFEST = PACK / "SAEE_SOFTWARE_COPYRIGHT_APPLICATION_MANIFEST_V1.json"
SOURCE_MANIFEST = PACK / "SAEE_SOFTWARE_COPYRIGHT_SOURCE_MANIFEST_V1.json"


def require(condition: bool, message: str) -> None:
    if not condition:
        raise SystemExit("SAEE_SOFTWARE_COPYRIGHT_APPLICATION_PACK_SMOKE: FAIL " + message)


def main() -> None:
    manifest = json.loads(MANIFEST.read_text(encoding="utf-8"))
    source = json.loads(SOURCE_MANIFEST.read_text(encoding="utf-8"))
    require(manifest["applicant"]["name"] == "山西游骑兵电子商务有限公司", "applicant")
    require(manifest["software"]["full_name"] == "SAEE智能体就绪评估软件", "software name")
    require(manifest["software"]["version"] == "V1.0", "software version")
    require(manifest["status"].startswith("hold_"), "fail-closed status")
    require(manifest["applicant"]["unified_social_credit_code"] == "91140802MA0GRJAX44", "credit code")
    require(manifest["applicant"]["contact_phone"] == "PRIVATE_LOCAL_VALUE_PRESENT", "phone privacy")
    require(manifest["applicant"]["mailing_address"] == "PRIVATE_LOCAL_VALUE_PRESENT", "address privacy")
    require(manifest["software"]["publication_status"] == "unpublished", "publication status")
    require(manifest["software"]["deposit_mode"] == "ordinary_deposit", "deposit mode")
    require(len(manifest["blocking_fields"]) == 3, "remaining gates")
    require(manifest["truth_boundary"]["owner_legal_fields_complete"] is True, "legal fields")
    require(manifest["truth_boundary"]["ownership_declaration_prepared"] is True, "declaration prepared")
    require(manifest["truth_boundary"]["ownership_declaration_signed_or_sealed"] is False, "signature overclaim")
    require(manifest["truth_boundary"]["application_submitted"] is False, "submission overclaim")
    require(manifest["truth_boundary"]["certificate_issued"] is False, "certificate overclaim")
    require(manifest["truth_boundary"]["production_ready"] is False, "production overclaim")
    require(source["complete_source_submitted_because_under_60_pages"] is True, "complete source rule")
    require(source["final_submission_ready"] is False, "source readiness overclaim")
    total = 0
    for item in source["files"]:
        path = ROOT / item["path"]
        require(path.is_file(), "missing source " + item["path"])
        require(hashlib.sha256(path.read_bytes()).hexdigest() == item["sha256"], "source hash drift " + item["path"])
        total += len(path.read_text(encoding="utf-8").splitlines())
    require(total == source["source_logical_line_count"], "source line count")
    for relative in manifest["generated_documents"]:
        path = ROOT / relative
        require(path.is_file() and path.stat().st_size > 10_000, "missing docx " + relative)
        document = Document(path)
        text_parts = [paragraph.text for paragraph in document.paragraphs]
        text_parts.extend(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        for section in document.sections:
            text_parts.extend(paragraph.text for paragraph in section.header.paragraphs)
            text_parts.extend(paragraph.text for paragraph in section.footer.paragraphs)
        text = "\n".join(text_parts)
        require("山西游骑兵电子商务有限公司" in text or "源程序鉴别材料" in text, "docx identity " + relative)
        require("970.jpg" not in text, "private license filename leaked into docx " + relative)
        require("本地文件" not in text, "internal evidence locator leaked into docx " + relative)
    print(
        "SAEE_SOFTWARE_COPYRIGHT_APPLICATION_PACK_SMOKE: PASS "
        f"source_files={len(source['files'])} logical_lines={total} "
        "documents=4 application_submitted=false certificate_issued=false"
    )


if __name__ == "__main__":
    main()
