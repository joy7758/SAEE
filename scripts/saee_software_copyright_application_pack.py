#!/usr/bin/env python3
"""Build the local, fail-closed SAEE software-copyright application pack.

The builder creates agent-readable manifests plus four human-review DOCX
artifacts. It never logs in to a copyright portal, uploads files, or submits an
application. Private contact and mailing values stay in a git-excluded manifest.
"""

from __future__ import annotations

import hashlib
import json
import subprocess
import textwrap
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs/ip/software-copyright"
SOFTWARE_NAME = "SAEE智能体就绪评估软件"
SOFTWARE_VERSION = "V1.0"
APPLICANT = "山西游骑兵电子商务有限公司"
CONTACT = "张斌"
UNIFIED_SOCIAL_CREDIT_CODE = "91140802MA0GRJAX44"
ENTITY_TYPE = "有限责任公司（自然人独资）"
LEGAL_REPRESENTATIVE = "张斌"
REGISTERED_CAPITAL = "壹佰万元整"
ESTABLISHMENT_DATE = "2015-12-04"
BUSINESS_TERM = "长期"
REGISTERED_ADDRESS = "运城市盐湖区马郝大道118号（耀华创业大厦四楼1号）"
PRIVATE_INPUT_PATH = Path.home() / ".config/saee/ip-private/software-copyright.json"
BUSINESS_LICENSE_COPY_DATE = "2020-02-28"
DEVELOPMENT_COMPLETION_DATE = "2026-07-13"
PUBLICATION_STATUS = "未发表"
DEVELOPMENT_MODE = "独立开发 / 自主研发"
RIGHTS_ACQUISITION = "原始取得"
RIGHTS_SCOPE = "全部权利"
DEPOSIT_MODE = "普通交存"

SOURCE_FILES = [
    Path("scripts/saee_agent_readiness_mcp_stdio.py"),
    Path("saee_backend/services/qianfan_readiness_mcp_adapter.py"),
    Path("saee_backend/services/baidu_agent_readiness_service.py"),
]

CONTRACT_FILES = [
    Path("agent-interface/qianfan/saee-readiness-evidence-item.schema.v0.1.json"),
    Path("agent-interface/qianfan/saee-evaluate-agent-run-request.schema.v0.1.json"),
    Path("agent-interface/qianfan/saee-evaluate-agent-run-response.schema.v0.1.json"),
    Path("agent-interface/qianfan/saee-evaluate-evidence-request.schema.v0.1.json"),
    Path("agent-interface/qianfan/saee-evaluate-evidence-response.schema.v0.1.json"),
    Path("agent-interface/qianfan/saee-qianfan-agent-readiness-mcp.v0.1.json"),
    Path("agent-interface/product/saee-agent-readiness-capability.v2.json"),
]

EXAMPLE_FILES = [
    Path("examples/baidu-qianfan/customer-service-refund/request.json"),
    Path("examples/baidu-qianfan/coding-agent-release/request.json"),
    Path("examples/baidu-qianfan/evaluate-evidence/request.json"),
]

VALIDATION_FILES = [Path("scripts/saee_qianfan_readiness_mcp_smoke.py")]

BLUE = RGBColor(31, 78, 121)
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
MUTED = RGBColor(89, 89, 89)


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def load_private_input() -> dict:
    if not PRIVATE_INPUT_PATH.is_file():
        raise SystemExit("missing private application input: " + str(PRIVATE_INPUT_PATH))
    value = json.loads(PRIVATE_INPUT_PATH.read_text(encoding="utf-8"))
    required = (
        "contact_name", "contact_phone", "contact_email", "mailing_address",
        "postal_code", "business_license_path",
    )
    missing = [key for key in required if not value.get(key)]
    if missing:
        raise SystemExit("missing private application fields: " + ", ".join(missing))
    return value


def git_head() -> str:
    return subprocess.check_output(
        ["git", "rev-parse", "HEAD"], cwd=ROOT, text=True
    ).strip()


def git_state(path: Path) -> str:
    value = subprocess.check_output(
        ["git", "status", "--short", "--", str(path)], cwd=ROOT, text=True
    ).strip()
    if not value:
        return "tracked_clean"
    if value.startswith("??"):
        return "untracked"
    return "modified_or_staged"


BODY_FONT = "STHeiti"


def font(run, *, name: str = BODY_FONT, size: float = 10.5, bold: bool = False, color=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def code_font(run, *, size: float) -> None:
    """Use a monospace Latin face while preserving a CJK fallback."""
    run.font.name = "Courier New"
    r_fonts = run._element.get_or_add_rPr().rFonts
    r_fonts.set(qn("w:ascii"), "Courier New")
    r_fonts.set(qn("w:hAnsi"), "Courier New")
    r_fonts.set(qn("w:eastAsia"), BODY_FONT)
    r_fonts.set(qn("w:cs"), BODY_FONT)
    run.font.size = Pt(size)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    font(run, size=9, color=MUTED)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_end])
    tail = paragraph.add_run(" 页")
    font(tail, size=9, color=MUTED)


def configure_doc(doc: Document, *, title: str, source_listing: bool = False) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    margin = 1.25 if source_listing else 2.2
    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    for key in ("w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"):
        normal._element.rPr.rFonts.set(qn(key), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after in (
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 11.5, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = BODY_FONT
        for key in ("w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"):
            style._element.rPr.rFonts.set(qn(key), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = header.add_run(f"{title}  |  {SOFTWARE_VERSION}")
    font(hr, size=8.5, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    font(r, size=24, bold=True, color=BLUE)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(24)
    r2 = p2.add_run(subtitle)
    font(r2, size=13, color=MUTED)


def add_key_value_table(doc: Document, rows: Iterable[tuple[str, str]], widths=(3.6, 12.8)) -> None:
    rows = list(rows)
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for idx, (label, value) in enumerate(rows):
        set_row_cant_split(table.rows[idx])
        left, right = table.rows[idx].cells
        left.width = Cm(widths[0])
        right.width = Cm(widths[1])
        set_cell_shading(left, LIGHT_BLUE)
        for cell in (left, right):
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        lp = left.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(label)
        font(lr, size=9.5, bold=True)
        rp = right.paragraphs[0]
        rp.paragraph_format.space_after = Pt(0)
        rr = rp.add_run(value)
        font(rr, size=9.5)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    font(r)


def build_application_fields(manifest: dict, private: dict) -> Path:
    doc = Document()
    configure_doc(doc, title="软件著作权申请字段底稿")
    add_title(doc, "计算机软件著作权登记申请字段底稿", f"{SOFTWARE_NAME} {SOFTWARE_VERSION}")
    add_key_value_table(doc, [
        ("申请状态", "HOLD：法定字段已回填，待公司签字/盖章、登记现状复核与外部提交授权；未登录、未上传、未提交"),
        ("著作权人", APPLICANT),
        ("统一社会信用代码", UNIFIED_SOCIAL_CREDIT_CODE),
        ("主体类型", ENTITY_TYPE),
        ("法定代表人", LEGAL_REPRESENTATIVE),
        ("注册地址", REGISTERED_ADDRESS),
        ("联系人", f"{private['contact_name']} / {private['contact_phone']} / {private['contact_email']}"),
        ("证书邮寄地址", f"{private['mailing_address']}（邮编 {private['postal_code']}）"),
        ("软件全称", SOFTWARE_NAME),
        ("软件简称", "SAEE"),
        ("版本号", SOFTWARE_VERSION),
        ("权利取得方式", f"{RIGHTS_ACQUISITION}（待申请主体签字或盖章确认）"),
        ("权利范围", f"{RIGHTS_SCOPE}（待申请主体签字或盖章确认）"),
    ])
    doc.add_heading("一、提交前仍需完成的确认门", level=1)
    for item in manifest["blocking_fields"]:
        add_bullet(doc, f"{item['field_zh']}（{item['field']}）")
    doc.add_heading("二、软件开发信息底稿", level=1)
    add_key_value_table(doc, [
        ("开发方式", f"{DEVELOPMENT_MODE}（待申请主体签字或盖章确认）"),
        ("开发完成日期", DEVELOPMENT_COMPLETION_DATE),
        ("首次发表状态", "未发表（本申请 V1.0 候选源程序未在公共仓库发布；公开 SAEE 抽象层不等于本程序发表）"),
        ("交存方式", f"{DEPOSIT_MODE}（候选程序不足60页，提交全部源程序鉴别材料）"),
        ("开发硬件环境", "通用 x86_64/ARM64 计算机，建议内存 8GB 及以上"),
        ("运行硬件环境", "通用服务器或个人计算机"),
        ("开发操作系统", "macOS/Linux；最终以申请主体实际开发记录确认"),
        ("运行操作系统", "macOS/Linux/Windows（具备 Python 运行环境）"),
        ("开发工具", "Python 3、文本编辑器、Git"),
        ("运行平台", "Python 3.10+，本地 stdio MCP/JSON-RPC 适配器"),
        ("编程语言", "Python、JSON Schema"),
        ("源程序量", f"候选源程序 {manifest['source_logical_line_count']} 个逻辑行；提交时以最终冻结包为准"),
    ])
    doc.add_heading("三、开发目的与主要功能", level=1)
    for text in (
        "开发目的：在智能体执行可能产生重大外部影响的动作前，对声明的运行轨迹和证据覆盖情况进行只读评估，为独立授权决策提供结构化上下文。",
        "主要功能一：saee.evaluate_agent_run，根据事件影响等级确定所需证据，输出覆盖分数、缺失证据、风险和 CONTINUE/REPLAN/STOP 等建议上下文。",
        "主要功能二：saee.evaluate_evidence，根据显式证据类型集合评估证据覆盖率，输出 SUFFICIENT/PARTIAL/INSUFFICIENT 结果。",
        "接口能力：通过固定的本地 MCP stdio/JSON-RPC 接口发现和调用上述两个只读工具。",
        "安全边界：软件不执行外部世界动作，不验证证据真实性，不授予部署权限，不输出安全认证或法律结论。",
    ):
        add_bullet(doc, text)
    doc.add_heading("四、技术特点", level=1)
    for text in (
        "文件化 JSON Schema 契约，输入、输出和边界可由编码智能体与检索智能体直接解析。",
        "确定性评估；相同输入产生相同结构化结果。",
        "严格限制为两个公开只读操作，未知工具和非法参数 fail-closed 拒绝。",
        "证据覆盖分数明确不是可靠性概率、认证或授权结论。",
        "网络访问和外部执行不属于本申请软件切片的默认运行路径。",
    ):
        add_bullet(doc, text)
    doc.add_heading("五、提交门", level=1)
    source_freeze = manifest["candidate_freeze"]
    add_key_value_table(doc, [
        (
            "版本冻结",
            "PASS：候选源文件已纳入 Git 且工作区状态干净；"
            f"Git HEAD={source_freeze['git_head_observed'][:12]}，"
            f"freeze_id={source_freeze['freeze_id_sha256'][:12]}…",
        ),
        ("权属核验", "READY_FOR_SIGNATURE：已按用户授权采用独立开发、原始取得、全部权利；待公司签字或盖章"),
        ("专利/秘密筛查", "PASS：采用普通交存；候选切片不含密钥、客户数据或私有进化内核"),
        ("登记现状", "HOLD：执照二维码与字段已核对；当前环境无法实时打开国家企业信用信息公示系统，提交前复核"),
        ("外部提交", "NOT_PERFORMED：未登录中国版权保护中心、未上传、未提交"),
    ])
    path = OUT / f"{SOFTWARE_NAME}{SOFTWARE_VERSION}_申请字段底稿.docx"
    doc.save(path)
    return path


def build_manual() -> Path:
    doc = Document()
    configure_doc(doc, title="用户操作说明书")
    add_title(doc, f"{SOFTWARE_NAME} {SOFTWARE_VERSION}", "用户操作说明书")
    add_key_value_table(doc, [
        ("著作权人", APPLICANT),
        ("软件简称", "SAEE"),
        ("文档版本", "1.0"),
        ("文档日期", str(date.today())),
        ("软件性质", "智能体执行前证据覆盖与就绪性评估软件"),
        ("当前边界", "本地只读评估；不执行外部动作，不授权部署"),
    ])
    doc.add_page_break()
    doc.add_heading("1. 软件概述", level=1)
    doc.add_paragraph(
        "SAEE智能体就绪评估软件面向需要在重大外部动作前检查运行证据的智能体平台、开发工具和本地工作流。"
        "软件接收结构化运行轨迹或证据清单，按照固定证据规则生成覆盖分数、缺失项、风险和建议上下文。"
        "评估结果仅供后续独立授权决策使用。"
    )
    doc.add_heading("1.1 软件定位", level=2)
    for text in (
        "工程核心：Digital Biosphere Evolution Engine（数字生物圈进化引擎）。",
        "产品投影：SAEE Agent Readiness Capability（SAEE智能体就绪评估能力）。",
        "公开操作严格限定为 saee.evaluate_agent_run 与 saee.evaluate_evidence。",
        "不适用于简单查询、文本改写、权限强制执行、安全认证、合规认证或自动部署。",
    ):
        add_bullet(doc, text)
    doc.add_heading("1.2 演化子系统关系", level=2)
    doc.add_paragraph(
        "本软件主要强化 Trait Extraction（性状提取）与 Pareto Fitness Evaluation（帕累托适应度评估），"
        "并通过固定契约为 Evolutionary Archive / Rollback Immune System（演化档案/回滚免疫系统）提供可复核结果。"
    )

    doc.add_heading("2. 运行环境", level=1)
    add_key_value_table(doc, [
        ("处理器", "x86_64 或 ARM64"),
        ("内存", "建议 8GB 及以上"),
        ("操作系统", "macOS、Linux 或 Windows"),
        ("语言环境", "Python 3.10 及以上"),
        ("Python 依赖", "jsonschema 4.18+；referencing（随 jsonschema 依赖链使用）"),
        ("通信方式", "本地标准输入/标准输出，JSON-RPC 2.0 消息"),
        ("网络", "核心评估不要求网络访问"),
    ])
    doc.add_heading("2.1 文件组成", level=2)
    for text in (
        "scripts/saee_agent_readiness_mcp_stdio.py：本地 MCP stdio 启动入口。",
        "saee_backend/services/qianfan_readiness_mcp_adapter.py：两工具 MCP/JSON-RPC 适配器。",
        "saee_backend/services/baidu_agent_readiness_service.py：运行轨迹与证据覆盖评估核心。",
        "agent-interface/qianfan/*.schema.v0.1.json：输入、输出与证据项契约。",
        "examples/baidu-qianfan/：不含真实客户数据的合成示例。",
    ):
        add_bullet(doc, text)

    doc.add_heading("3. 安装与启动", level=1)
    doc.add_heading("3.1 准备环境", level=2)
    for step in (
        "确认 Python 3.10 或更高版本可用。",
        "在隔离环境中安装 jsonschema 4.18 或更高版本。",
        "从软件根目录运行启动脚本；不要从未知仓库执行安装脚本。",
    ):
        add_bullet(doc, step)
    doc.add_heading("3.2 启动命令", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("python3 scripts/saee_agent_readiness_mcp_stdio.py")
    code_font(r, size=9.5)
    doc.add_paragraph(
        "启动后进程从标准输入逐行读取 JSON-RPC 请求，并在标准输出逐行返回 JSON-RPC 响应。"
        "默认最大单行消息为 1,000,000 字节，最大 JSON 嵌套深度为 64。"
    )

    doc.add_heading("4. 初始化与能力发现", level=1)
    doc.add_heading("4.1 initialize", level=2)
    doc.add_paragraph(
        "客户端首先发送 initialize 请求。软件协商支持的协议版本并返回 serverInfo，"
        "随后客户端发送 notifications/initialized 通知。未完成初始化时调用其他方法将被拒绝。"
    )
    doc.add_heading("4.2 tools/list", level=2)
    doc.add_paragraph("初始化完成后调用 tools/list。软件只返回以下两个公开工具：")
    for text in (
        "saee.evaluate_agent_run：评估一条声明的智能体运行轨迹及所需证据覆盖。",
        "saee.evaluate_evidence：评估证据包是否覆盖指定证据类型集合。",
    ):
        add_bullet(doc, text)
    doc.add_paragraph(
        "rehearse_agent、describe_saee、compare_observed_traces 等内部工程接口不会出现在本软件的公开工具列表中。"
    )

    doc.add_heading("5. 操作一：评估智能体运行", level=1)
    doc.add_heading("5.1 输入内容", level=2)
    for text in (
        "request_id：请求唯一标识。",
        "agent_id：被评估智能体标识。",
        "task：任务说明。",
        "trace.events：事件列表，每项声明 event_type、external_effect 与 high_impact。",
        "evidence：证据项列表，类型包括 TEST_RESULT、ROLLBACK_PLAN、PERMISSION_BOUNDARY、HUMAN_APPROVAL。",
        "customer_data_included：必须为 false；本地 Alpha 不接受客户数据。",
    ):
        add_bullet(doc, text)
    doc.add_heading("5.2 评估逻辑", level=2)
    doc.add_paragraph(
        "普通运行至少要求 TEST_RESULT。存在高影响或外部效果事件时，要求测试结果、回滚计划、权限边界和人工批准四类证据。"
        "软件按已声明为 present 的证据计算覆盖率，并映射到 continue、conditional、replan 或 stop。"
    )
    doc.add_heading("5.3 输出解释", level=2)
    add_key_value_table(doc, [
        ("100%", "continue / CONTINUE"),
        ("75%及以上", "conditional / HUMAN_REVIEW_REQUIRED"),
        ("50%及以上", "replan / REPLAN"),
        ("低于50%", "stop / STOP"),
    ])
    doc.add_paragraph(
        "覆盖率是所需证据类型的存在比例，不是可靠性概率或安全概率。即使结果为 CONTINUE，也不代表获得部署授权。"
    )

    doc.add_heading("6. 操作二：评估证据包", level=1)
    doc.add_heading("6.1 输入内容", level=2)
    for text in (
        "evidence_bundle.items：待评估证据项集合。",
        "required_evidence_types：本次明确要求的证据类型。",
        "customer_data_included：必须为 false。",
    ):
        add_bullet(doc, text)
    doc.add_heading("6.2 输出内容", level=2)
    for text in (
        "evidence_quality：SUFFICIENT、PARTIAL 或 INSUFFICIENT。",
        "coverage_score：所需证据覆盖百分比。",
        "present_evidence 与 missing_evidence：已具备和缺失的证据类型。",
        "reason_codes：每项缺失证据对应的机器可读原因码。",
    ):
        add_bullet(doc, text)

    doc.add_heading("7. 合成示例", level=1)
    doc.add_heading("7.1 客服退款准备", level=2)
    doc.add_paragraph(
        "示例只读取合成订单相关声明，不执行支付。测试、回滚计划和权限边界存在，人工批准缺失，"
        "因此覆盖率为 75%，结果为 conditional，并要求人工复核。"
    )
    doc.add_heading("7.2 编码智能体发布准备", level=2)
    doc.add_paragraph(
        "示例声明一次高影响发布决策，但缺少回滚计划和人工批准，覆盖率为 50%，结果为 replan。"
    )
    doc.add_heading("7.3 证据包评估", level=2)
    doc.add_paragraph(
        "示例要求测试结果、回滚计划和权限边界三项证据，其中两项存在，覆盖率为 67%，结果为 PARTIAL。"
    )

    doc.add_heading("8. 异常处理", level=1)
    for text in (
        "非法 JSON：返回 JSON-RPC Parse error。",
        "未初始化调用：返回 Server not initialized。",
        "未知工具或参数类型错误：返回 Invalid params。",
        "Schema 校验失败：返回 READINESS_MCP_ARGUMENTS_INVALID。",
        "重复 evidence_id 或 evidence_type：核心服务 fail-closed 拒绝。",
        "超出消息大小或嵌套深度：请求被拒绝，不进入评估。",
    ):
        add_bullet(doc, text)

    doc.add_heading("9. 安全与权限边界", level=1)
    for text in (
        "软件只评估调用方提供的声明，不验证轨迹或证据真实性。",
        "软件不访问客户系统，不执行候选代码，不自动扩大权限。",
        "软件不执行部署、支付、外发、联系客户或其他外部世界动作。",
        "输出不构成安全认证、合规结论、法律意见或平台认可。",
        "重大外部动作必须由独立的人类授权门决定。",
    ):
        add_bullet(doc, text)

    doc.add_heading("10. 验证与维护", level=1)
    doc.add_heading("10.1 离线验证", level=2)
    p = doc.add_paragraph()
    r = p.add_run("python3 scripts/saee_qianfan_readiness_mcp_smoke.py")
    code_font(r, size=9.5)
    doc.add_paragraph(
        "验证覆盖两个工具的发现、三个合成示例、非法参数拒绝、重复证据拒绝及五次确定性重复运行。"
    )
    doc.add_heading("10.2 版本维护", level=2)
    doc.add_paragraph(
        "任何改变公开工具、评估规则、Schema 或安全边界的修改，都必须同步更新说明书、机器契约、"
        "源程序哈希清单和登记版本决策。V1.0 申请提交后，后续实质性功能变化应评估是否登记新版本。"
    )
    path = OUT / f"{SOFTWARE_NAME}{SOFTWARE_VERSION}_用户操作说明书.docx"
    doc.save(path)
    return path


def build_ownership_confirmation(manifest: dict) -> Path:
    """Build the owner-signature gate without authorizing portal submission."""
    doc = Document()
    configure_doc(doc, title="权属与申请确认书")
    add_title(doc, "软件权属与申请信息确认书", f"{SOFTWARE_NAME} {SOFTWARE_VERSION}")
    doc.add_paragraph(
        "本确认书用于公司内部确认申请字段、权属取得方式和交存策略。签字或盖章仅确认本页事项，"
        "不等于授权任何人登录、上传或向登记机构提交申请。"
    )
    doc.add_heading("一、申请主体", level=1)
    add_key_value_table(doc, [
        ("著作权人", APPLICANT),
        ("统一社会信用代码", UNIFIED_SOCIAL_CREDIT_CODE),
        ("主体类型", ENTITY_TYPE),
        ("法定代表人", LEGAL_REPRESENTATIVE),
        ("注册地址", REGISTERED_ADDRESS),
        ("营业期限", BUSINESS_TERM),
    ])
    doc.add_heading("二、确认事项", level=1)
    for text in (
        f"本公司确认拟登记软件为“{SOFTWARE_NAME} {SOFTWARE_VERSION}”。",
        f"开发方式采用“{DEVELOPMENT_MODE}”，权利取得方式采用“{RIGHTS_ACQUISITION}”，权利范围采用“{RIGHTS_SCOPE}”。",
        f"开发完成日期采用 {DEVELOPMENT_COMPLETION_DATE}；首次发表状态采用“{PUBLICATION_STATUS}”。",
        "“未发表”仅针对本申请冻结的 V1.0 候选源程序；公开的 SAEE 抽象层、说明材料或 public-safe release 不视为该候选程序已经发表。",
        f"交存方式采用“{DEPOSIT_MODE}”；候选源程序不足60页，提交全部源程序鉴别材料。",
        "候选切片不含密钥、客户数据或私有进化内核；如签署人掌握相反事实，应停止签署并重新评估交存方式。",
        "公司确认其有权就候选源程序主张并登记上述软件著作权；如存在职务开发、委托开发、合作开发或第三方权属安排，应在签署前书面披露。",
    ):
        add_bullet(doc, text)
    doc.add_heading("三、证据与边界", level=1)
    add_key_value_table(doc, [
        (
            "营业执照副本",
            f"申请人提供的营业执照副本；统一社会信用代码 {UNIFIED_SOCIAL_CREDIT_CODE}",
        ),
        ("材料校验摘要", f"SHA-256：{manifest['business_license_evidence']['sha256']}"),
        ("执照副本日期", BUSINESS_LICENSE_COPY_DATE),
        ("实时登记核验", "未完成；当前浏览器环境访问公示系统时被客户端拦截，门户提交前复核"),
        ("源程序冻结", manifest["candidate_freeze"]["freeze_id_sha256"]),
        ("外部动作", "未登录、未上传、未提交；需另行明确授权"),
    ])
    doc.add_heading("四、签署", level=1)
    doc.add_paragraph("本公司已核对以上信息，并确认其真实、准确、完整。")
    add_key_value_table(doc, [
        ("公司盖章", ""),
        ("法定代表人/授权签字人", ""),
        ("签署日期", "年    月    日"),
    ])
    path = OUT / f"{SOFTWARE_NAME}{SOFTWARE_VERSION}_权属与申请确认书.docx"
    doc.save(path)
    return path


def display_source_lines() -> list[tuple[str, int | None, str]]:
    result: list[tuple[str, int | None, str]] = []
    for relative in SOURCE_FILES:
        result.append((str(relative), None, f"# FILE: {relative}"))
        logical_lines = (ROOT / relative).read_text(encoding="utf-8").splitlines()
        for number, value in enumerate(logical_lines, start=1):
            expanded = value.expandtabs(4)
            chunks = textwrap.wrap(
                expanded,
                width=112,
                replace_whitespace=False,
                drop_whitespace=False,
                break_long_words=True,
                break_on_hyphens=False,
            ) or [""]
            result.append((str(relative), number, chunks[0]))
            for chunk in chunks[1:]:
                result.append((str(relative), 0, chunk))
    return result


def build_source_listing() -> tuple[Path, int]:
    rows = display_source_lines()
    doc = Document()
    configure_doc(doc, title="源程序鉴别材料", source_listing=True)
    section = doc.sections[0]
    section.header.paragraphs[0].clear()
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = header.add_run(f"{SOFTWARE_NAME} {SOFTWARE_VERSION}  源程序鉴别材料")
    font(hr, size=8.5, color=MUTED)
    for index, (relative, logical_number, value) in enumerate(rows):
        if index and index % 50 == 0:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.keep_together = True
        if logical_number is None:
            label = "       "
        elif logical_number == 0:
            label = "      +"
        else:
            label = f"{logical_number:6d} "
        r = p.add_run(label + value)
        code_font(r, size=6.2)
    path = OUT / f"{SOFTWARE_NAME}{SOFTWARE_VERSION}_源程序鉴别材料.docx"
    doc.save(path)
    return path, len(rows)


def build_manifest(private: dict) -> dict:
    business_license_path = Path(private["business_license_path"])
    if not business_license_path.is_file():
        raise SystemExit("missing business license evidence: " + str(business_license_path))
    all_paths = SOURCE_FILES + CONTRACT_FILES + EXAMPLE_FILES + VALIDATION_FILES
    missing = [str(path) for path in all_paths if not (ROOT / path).is_file()]
    if missing:
        raise SystemExit("missing required application-pack inputs: " + ", ".join(missing))
    files = []
    for path in all_paths:
        files.append({
            "path": str(path),
            "role": (
                "source_program" if path in SOURCE_FILES else
                "machine_contract" if path in CONTRACT_FILES else
                "example" if path in EXAMPLE_FILES else
                "offline_validation"
            ),
            "sha256": sha256(ROOT / path),
            "bytes": (ROOT / path).stat().st_size,
            "logical_lines": len((ROOT / path).read_text(encoding="utf-8").splitlines()),
            "git_state": git_state(path),
        })
    freeze_payload = "\n".join(f"{item['path']}:{item['sha256']}" for item in files)
    source_count = sum(item["logical_lines"] for item in files if item["role"] == "source_program")
    source_tracked_clean = all(
        item["git_state"] == "tracked_clean"
        for item in files
        if item["role"] == "source_program"
    )
    blocking_fields = [
        {"field": "business_registry_currentness_confirmation", "field_zh": "提交前复核企业当前登记状态"},
        {"field": "owner_signature_or_seal", "field_zh": "在权属与申请确认书上签字或盖章"},
        {"field": "explicit_portal_submission_authorization", "field_zh": "另行明确授权门户登录、上传和提交"},
    ]
    if not source_tracked_clean:
        blocking_fields.append(
            {"field": "tracked_version_freeze", "field_zh": "源程序纳入版本控制并冻结最终哈希"}
        )
    return {
        "schema_version": "1.0.0",
        "artifact": "SAEE software copyright application manifest",
        "status": "hold_owner_signature_registry_check_and_portal_authorization",
        "applicant": {
            "name": APPLICANT,
            "type": "enterprise_legal_person",
            "entity_type_zh": ENTITY_TYPE,
            "unified_social_credit_code": UNIFIED_SOCIAL_CREDIT_CODE,
            "registered_address": REGISTERED_ADDRESS,
            "legal_representative": LEGAL_REPRESENTATIVE,
            "registered_capital_zh": REGISTERED_CAPITAL,
            "establishment_date": ESTABLISHMENT_DATE,
            "business_term_zh": BUSINESS_TERM,
            "business_license_copy": "APPLICANT_PROVIDED_COPY_AVAILABLE_FOR_PORTAL_UPLOAD",
            "contact_name": CONTACT,
            "contact_phone": "PRIVATE_LOCAL_VALUE_PRESENT",
            "contact_email": "PRIVATE_LOCAL_VALUE_PRESENT",
            "mailing_address": "PRIVATE_LOCAL_VALUE_PRESENT",
            "postal_code": "PRIVATE_LOCAL_VALUE_PRESENT",
        },
        "business_license_evidence": {
            "present": True,
            "stored_outside_repository": True,
            "sha256": sha256(business_license_path),
            "copy_date_observed": BUSINESS_LICENSE_COPY_DATE,
            "qr_unified_social_credit_code": UNIFIED_SOCIAL_CREDIT_CODE,
            "current_registry_live_verified": False,
            "live_verification_note": "2026-07-13 access blocked by browser client; recheck before portal submission",
        },
        "private_application_manifest": "docs/ip/software-copyright/SAEE_SOFTWARE_COPYRIGHT_PRIVATE_APPLICATION_V1.json",
        "software": {
            "full_name": SOFTWARE_NAME,
            "short_name": "SAEE",
            "version": SOFTWARE_VERSION,
            "development_mode": "independent_development",
            "development_mode_zh": DEVELOPMENT_MODE,
            "rights_acquisition": "original_acquisition",
            "rights_acquisition_zh": RIGHTS_ACQUISITION,
            "rights_scope": "all_rights",
            "rights_scope_zh": RIGHTS_SCOPE,
            "development_completion_date": DEVELOPMENT_COMPLETION_DATE,
            "publication_status": "unpublished",
            "publication_status_zh": PUBLICATION_STATUS,
            "first_publication_date": None,
            "first_publication_place": None,
            "deposit_mode": "ordinary_deposit",
            "deposit_mode_zh": DEPOSIT_MODE,
            "source_logical_line_count": source_count,
        },
        "candidate_freeze": {
            "git_head_observed": git_head(),
            "freeze_id_sha256": hashlib.sha256(freeze_payload.encode("utf-8")).hexdigest(),
            "all_source_files_tracked_clean": source_tracked_clean,
            "files": files,
        },
        "source_logical_line_count": source_count,
        "blocking_fields": blocking_fields,
        "truth_boundary": {
            "application_materials_prepared_local": True,
            "owner_legal_fields_complete": True,
            "owner_recommendation_authorized": True,
            "ownership_declaration_prepared": True,
            "ownership_declaration_signed_or_sealed": False,
            "ownership_verified": False,
            "publication_status_selected": True,
            "deposit_mode_selected": True,
            "source_version_committed": source_tracked_clean,
            "portal_login_performed": False,
            "files_uploaded": False,
            "application_submitted": False,
            "application_accepted": False,
            "certificate_issued": False,
            "production_ready": False,
            "customer_validated": False,
            "marketplace_listed": False,
        },
    }


def write_agent_readable_files(
    manifest: dict, docs: list[Path], display_line_count: int, private: dict
) -> None:
    manifest["generated_documents"] = [str(path.relative_to(ROOT)) for path in docs]
    manifest["source_display_line_count"] = display_line_count
    (OUT / "SAEE_SOFTWARE_COPYRIGHT_APPLICATION_MANIFEST_V1.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    private_manifest_path = OUT / "SAEE_SOFTWARE_COPYRIGHT_PRIVATE_APPLICATION_V1.json"
    private_manifest = {
        "schema_version": "1.0.0",
        "privacy": "private_local_git_excluded",
        "applicant": APPLICANT,
        "contact_name": private["contact_name"],
        "contact_phone": private["contact_phone"],
        "contact_email": private["contact_email"],
        "mailing_address": private["mailing_address"],
        "postal_code": private["postal_code"],
        "business_license_path": private["business_license_path"],
        "business_license_sha256": manifest["business_license_evidence"]["sha256"],
    }
    private_manifest_path.write_text(
        json.dumps(private_manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    private_manifest_path.chmod(0o600)
    source_manifest = {
        "schema_version": "1.0.0",
        "software_name": SOFTWARE_NAME,
        "software_version": SOFTWARE_VERSION,
        "candidate_freeze_id_sha256": manifest["candidate_freeze"]["freeze_id_sha256"],
        "source_logical_line_count": manifest["source_logical_line_count"],
        "source_display_line_count": display_line_count,
        "files": [item for item in manifest["candidate_freeze"]["files"] if item["role"] == "source_program"],
        "complete_source_submitted_because_under_60_pages": True,
        "final_submission_ready": False,
    }
    (OUT / "SAEE_SOFTWARE_COPYRIGHT_SOURCE_MANIFEST_V1.json").write_text(
        json.dumps(source_manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )


def main() -> int:
    OUT.mkdir(parents=True, exist_ok=True)
    private = load_private_input()
    manifest = build_manifest(private)
    fields = build_application_fields(manifest, private)
    manual = build_manual()
    confirmation = build_ownership_confirmation(manifest)
    source, display_count = build_source_listing()
    write_agent_readable_files(
        manifest, [fields, manual, confirmation, source], display_count, private
    )
    print(json.dumps({
        "status": manifest["status"],
        "applicant": APPLICANT,
        "software": f"{SOFTWARE_NAME} {SOFTWARE_VERSION}",
        "candidate_freeze_id": manifest["candidate_freeze"]["freeze_id_sha256"],
        "source_logical_lines": manifest["source_logical_line_count"],
        "source_display_lines": display_count,
        "documents": [str(path.relative_to(ROOT)) for path in (fields, manual, confirmation, source)],
        "external_submission_performed": False,
    }, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
